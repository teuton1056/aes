import aes_statistics
import docx
import csv
import xml.etree.ElementTree as ET
import html

class orthographic_info:

	def __init__(self,uid,tid,line,hiero_data):
		self.uid = uid
		self.tid = tid
		self.name = html.unescape(self.lookup_tid(self.tid))
		self.corpus, self.text_name = self.name.split(':')[0:2]
		self.line = line
		self.hiero_data = hiero_data

	def lookup_tid(self,tid):
		with open('concordance_name_text_id.csv','r',newline='',encoding='utf-8') as fp:
			reader = csv.reader(fp,delimiter='\t')
			for line in reader:
				if line[1] == tid:
					return line[0]
		return 'Error, no corpus:Error, the Text id was not found'

	def return_hiero(self):
		return self.hiero_data 
	
	def return_info(self):
		return f'({self.text_name}, {self.line})'

	def return_text_line(self):
		return f'{self.text_name}, {self.line}'

				
		

def find_spellings(lemma_id,mode='hiero_unicode'): 
	# acceptable modes are 'hiero_unicode','hiero' and 'hiero_inventar' 
	#the last is not possible if generating a report unless using the "invertory" report option.
	data = aes_statistics.load_data()
	total_spellings = 0
	results = {}
	for sentence in data:
		for token in sentence['token']:
			if 'lemmaID' in token.keys():
				if lemma_id == token['lemmaID']:
					if mode in token.keys():
						total_spellings += 1
						#print()
						if html.unescape(token[mode]) in results.keys():
							results[html.unescape(token[mode])][0] += 1
							results[html.unescape(token[mode])][1].append(orthographic_info(token['_id'],sentence['text'],token['lineCount'],html.unescape(token[mode])))
						else:
							results[html.unescape(token[mode])] = [1,[orthographic_info(token['_id'],sentence['text'],token['lineCount'],html.unescape(token[mode]))]]
	return total_spellings, results 

def lookup_word(word_id):
	root = ET.parse('dictionary.xml').getroot()
	ns = {'d':'http://www.tei-c.org/ns/1.0','xml':'http://www.w3.org/XML/1998/namespace'}
	entries = root.findall('.//d:entry',ns)
	for entry in entries:
		if entry.attrib[f'{{{ns["xml"]}}}id'] == f'tla{word_id}':
			lemma = entry.find('.//d:orth',ns).text
			bib = entry.find('.//d:bibl',ns).text
			return lemma, bib


def generate_report(word_id,mode='hiero_unicode',inventory=False):
	total, spellings = find_spellings(word_id,mode)
	print(f'A Total of {total} spellings were analyzed with the following results:')
	lemma, bib = lookup_word(word_id)
	document = docx.Document()
	document.add_heading(f'Orthographic Summary of {lemma}')
	document.add_paragraph(f"This orthographic summary is based on the AES databank. The queried word was {lemma} with id {word_id}. A total of {total} spellings were found.")
	document.add_paragraph(f"{bib}")
	keys = list(spellings.keys())
	table = document.add_table(rows=len(spellings.keys()),cols=3)
	j = 0
	while j < len(spellings.keys()):
		# okay so each entry in spellings looks like this [ total_number
		# [orthographic_info objects]
		# ]
		total = spellings[keys[j]][0]
		orths = [info.return_text_line() for info in spellings[keys[j]][1]]
		table.rows[j].cells[0].text = keys[j]
		table.rows[j].cells[1].text = str(total)
		table.rows[j].cells[2].text = '\n'.join(orths)
		j += 1
	document.save(f'{lemma}, {word_id}.docx')


def main():
	total, results = find_spellings('78030',mode='hiero')
	print(f'A Total of {total} spellings were analyzed with the following results:')
	print(results)

if __name__ == '__main__':
	generate_report('854561')